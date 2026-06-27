#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把器件清单(BOM)生成为横向排版的 Word 文档。

输出: 07_project_result_analysis/器件清单_BOM.docx
依赖: python-docx (仓库已用于 build_report_docx.py)
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "07_project_result_analysis" / "器件清单_BOM.docx"

FONT = "Microsoft YaHei"
HEAD_FILL = "2E5C8A"      # 表头深蓝
HEAD_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA = "EEF3F8"         # 隔行浅蓝
GRAY = RGBColor(0x60, 0x60, 0x60)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_font(run, size=Pt(9), bold=False, color=BLACK, name=FONT):
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def shade(cell, fill):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def cell_text(cell, text, size=Pt(8.5), bold=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)
    set_font(run, size=Pt(8), color=GRAY)


# 列定义: (标题, 宽度英寸)
COLS = [
    ("#", 0.3), ("名称", 1.05), ("型号 / 订货号", 1.3), ("厂商", 0.75),
    ("关键规格 / 接口", 2.2), ("数量", 0.5), ("单价(约)", 0.9),
    ("小计(约)", 0.9), ("用途", 1.1),
]

ROWS = [
    ["1", "RX 主板（本设计载体）", "Genesys ZU-5EV / 410-383-5EV", "Digilent",
     "Zynq UltraScale+ XCZU5EV；DDR4；2×Pcam(2-lane)口；FMC；HDMI/10G", "1",
     "¥1.6–1.8万", "¥1.6–1.8万", "跑 RTL（CSI-2 RX+重采集），PS DDR4 帧缓冲"],
    ["2", "可控 MIPI 发送端", "Zybo Z7-20 / 410-351-20", "Digilent",
     "Zynq-7020 XC7Z020-1CLG400C；Pcam 口", "1",
     "¥2.8–3.2k", "¥2.8–3.2k", "配成可控 MIPI TX，演示重采集闭环"],
    ["3", "2-lane 相机", "Pcam 5C / 410-358", "Digilent",
     "OV5640 5MP；2-lane MIPI CSI-2；Pcam 口直插", "1",
     "¥0.5–0.7k", "¥0.5–0.7k", "接 Genesys Pcam 口，验 2-lane 真实采集"],
    ["4", "4-lane 相机", "LI-IMX274MIPI-FMC", "Leopard Imaging",
     "Sony IMX274 8.5MP；4-lane MIPI CSI-2；FMC 接口", "1",
     "¥2.5–4k（询价）", "¥2.5–4k", "接 Genesys FMC 口，验 4-lane 真实采集"],
    ["5", "电源（主板）", "Genesys ZU 原配 12V 适配器", "Digilent",
     "12V（随板）", "1", "随板", "—", "Genesys 供电"],
    ["6", "电源（Zybo）", "5V/2.5A 5.5×2.1mm 中心正", "通用",
     "桶形插头", "1", "¥30–50", "¥30–50", "Zybo 供电"],
    ["7", "micro-USB 线", "USB-A ↔ micro-USB", "通用",
     "JTAG/UART 编程", "2", "¥20", "¥40", "两块板各一根"],
    ["8", "microSD 卡", "16GB Class10", "通用",
     "—", "2", "¥30", "¥60", "PS 端启动镜像"],
    ["9", "反向通道连线", "杜邦线（母-母）", "通用",
     "GPIO/I²C", "1套", "¥15", "¥15", "两板间传重发请求（反向通道）"],
    ["10", "（可选）第 2 个 2-lane 相机", "Pcam 5C / 410-358", "Digilent",
     "同 #3", "(1)", "¥0.5–0.7k", "(+¥0.5–0.7k)", "多相机/VC 演示"],
]

NUM_COLS = {0, 5, 6, 7}  # 居中列


def main():
    doc = Document()
    # 默认字体
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11), Inches(8.5)
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin = sec.bottom_margin = Inches(0.7)

    # 页脚
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("MIPI CSI-2 采集系统 · 板级测试器件清单 BOM      第 ")
    set_font(r, size=Pt(8), color=GRAY)
    add_field(fp, "PAGE")
    r = fp.add_run(" 页")
    set_font(r, size=Pt(8), color=GRAY)

    # 标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run("板级测试器件清单（BOM）")
    set_font(run, size=Pt(17), bold=True)

    sub = doc.add_paragraph()
    run = sub.add_run("MIPI CSI-2 图像采集系统 · 平台限定 Xilinx（Zynq / UltraScale+）· 预算 ¥5 万内，实花约 ¥2.5–3 万")
    set_font(run, size=Pt(9.5), color=GRAY)

    # 表格
    table = doc.add_table(rows=1 + len(ROWS), cols=len(COLS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.style = "Table Grid"

    # 表头
    hdr = table.rows[0].cells
    for j, (title, _) in enumerate(COLS):
        shade(hdr[j], HEAD_FILL)
        cell_text(hdr[j], title, size=Pt(9), bold=True, color=HEAD_TEXT,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # 数据行
    for i, row in enumerate(ROWS):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            if (i % 2) == 1:
                shade(cells[j], ZEBRA)
            align = WD_ALIGN_PARAGRAPH.CENTER if j in NUM_COLS else WD_ALIGN_PARAGRAPH.LEFT
            bold = (j == 0)
            cell_text(cells[j], val, size=Pt(8.3), bold=bold, align=align)

    # 固定列宽
    for j, (_, w) in enumerate(COLS):
        for row in table.rows:
            row.cells[j].width = Inches(w)

    # 合计
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("合计（不含可选 #10）≈ ¥22,000–26,000　｜　¥5 万预算内，余量 ¥2.4–2.8 万。")
    set_font(run, size=Pt(11), bold=True)

    # 连接关系
    doc.add_paragraph().add_run("")
    h = doc.add_paragraph(); run = h.add_run("连接关系"); set_font(run, size=Pt(12), bold=True)
    for line in [
        "Genesys ZU-5EV（RX，你的电路 + PS DDR4）：Pcam 口 ← Pcam 5C（2-lane 采集）；FMC 口 ← LI-IMX274MIPI-FMC（4-lane 采集）。",
        "Zybo Z7-20（可控 MIPI 发送端）── MIPI 数据 ──▶ Genesys；反向通道（杜邦线 GPIO/I²C）回传“重发请求”，演示行级重采集闭环。",
        "两个相机都接 RX 主板；Zybo 不接相机（它本身就是数据源）。重采集闭环用 Zybo 可控源演示，与两相机是相互独立的两类验证。",
    ]:
        bp = doc.add_paragraph(style="List Bullet")
        set_font(bp.add_run(line), size=Pt(10))

    # 采购要点
    h = doc.add_paragraph(); run = h.add_run("采购要点 / 注意"); set_font(run, size=Pt(12), bold=True)
    for line in [
        "认准订货号：410-383-5EV（主板）、410-351-20（Zybo，务必 -20 不是 -10）、410-358（Pcam 5C）。",
        "4-lane 相机需询价：LI-IMX274MIPI-FMC 为 Leopard Imaging B2B 产品（support@leopardimaging.com 或 Future Electronics）；备选 e-con Systems。",
        "4-lane 集成有工作量：该 FMC 相机官方主要在 ZCU102/104 验证，接 Genesys ZU 的 FMC 电气可行但非官方验证组合，需自写约束 + 配 4-lane D-PHY IP，非即插即用。",
        "不采购：MIPI 协议分析仪等高价实验室设备（远超本项目需求）。",
    ]:
        np = doc.add_paragraph(style="List Number")
        set_font(np.add_run(line), size=Pt(10))

    doc.save(str(OUT))
    print(f"written: {OUT}")


if __name__ == "__main__":
    main()
