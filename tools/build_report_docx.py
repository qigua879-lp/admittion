from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT_MD = ROOT / "MIPI CSI-2数字逻辑RX端设计及仿真验证说明.md"
OUTPUT_DOCX = ROOT / "MIPI CSI-2数字逻辑RX端设计及仿真验证说明.docx"


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
ORDERED_RE = re.compile(r"^(\d+)\.\s*(.*)$")

BODY_FONT = "Microsoft YaHei"
BODY_SIZE = Pt(10.5)
TABLE_BODY_SIZE = Pt(8.2)
TABLE_HEAD_SIZE = Pt(9.2)


def set_run_font(run, name=BODY_FONT, size=BODY_SIZE, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_margins(cell, top=45, start=85, bottom=45, end=85):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    ppr.append(keep)


def set_paragraph_format(paragraph, before=0, after=4, line=1.18):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def set_first_line_indent(paragraph, chars=2.0):
    paragraph.paragraph_format.first_line_indent = Pt(BODY_SIZE.pt * chars)


def add_markdown_runs(paragraph, text, default_font=BODY_FONT, size=BODY_SIZE, color=None):
    clean_text = text.replace("`", "")
    run = paragraph.add_run(clean_text)
    set_run_font(run, name=default_font, size=size, color=color)


def parse_table(lines):
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip().replace("`", "") for c in line.strip().strip("|").split("|")]
        if cells and all(set(c) <= {"-", ":"} for c in cells):
            continue
        rows.append(cells)
    return rows


def set_table_widths(table, headers):
    widths_map = {
        ("预期内容", "当前完成情况", "对比结论"): [Cm(4.0), Cm(8.0), Cm(3.5)],
        ("信号", "方向", "位宽", "时钟域", "说明"): [Cm(3.0), Cm(1.5), Cm(1.2), Cm(2.0), Cm(7.0)],
        ("信号组", "方向", "时钟域", "说明"): [Cm(3.0), Cm(2.0), Cm(2.0), Cm(7.0)],
        ("项目", "结果"): [Cm(5.0), Cm(4.0)],
        ("tb名称", "验证模块", "验证功能", "验证结果"): [Cm(4.2), Cm(3.2), Cm(4.0), Cm(6.4)],
        ("格式", "init_to_frame", "frame_to_first_pixel", "frame_to_end", "first_to_last_pixel", "pixel_valid_cycles", "闭环结果"): [Cm(2.0), Cm(2.0), Cm(2.6), Cm(2.1), Cm(2.4), Cm(2.3), Cm(2.0)],
        ("lane 数", "对应用例", "exp pixels", "act pixels", "frames", "结果说明"): [Cm(1.5), Cm(5.0), Cm(2.0), Cm(2.0), Cm(1.6), Cm(4.2)],
        ("用例", "关注点", "关键结果", "结论"): [Cm(4.8), Cm(3.6), Cm(6.2), Cm(3.8)],
    }
    key = tuple(headers)
    widths = widths_map.get(key)
    if widths is None:
        count = len(headers)
        widths = [Cm(16.0 / count)] * count
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = width


def build_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, rows[0])

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or len(value) < 14) else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(para, before=0, after=0, line=1.05)
            run = para.add_run(value)
            if r_idx == 0:
                set_run_font(run, size=TABLE_HEAD_SIZE, bold=True)
                shade_cell(cell, "D9E2F3")
            else:
                set_run_font(run, size=TABLE_BODY_SIZE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        if r_idx == 0:
            set_repeat_table_header(row)


def insert_image(doc, image_path, alt_text):
    image_file = (ROOT / image_path).resolve()
    if not image_file.exists():
        p = doc.add_paragraph()
        add_markdown_runs(p, f"[图片缺失] {alt_text}: {image_path}", size=Pt(10))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_file), width=Inches(5.85))
    set_paragraph_format(p, before=3, after=1, line=1.0)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(alt_text)
    set_run_font(cap_run, size=Pt(9), color=RGBColor(90, 90, 90))
    cap_run.italic = True
    set_paragraph_format(cap, before=0, after=5, line=1.0)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.05)
    section.right_margin = Cm(2.05)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
    normal.font.size = BODY_SIZE

    title = styles["Title"]
    title.font.name = BODY_FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    title._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
    title.font.size = Pt(18)
    title.font.bold = True

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)]:
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)


def build_docx():
    lines = INPUT_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        heading = HEADING_RE.match(line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_run_font(run, size=Pt(18), bold=True)
                set_paragraph_format(p, before=0, after=12, line=1.0)
            else:
                p = doc.add_paragraph(style=f"Heading {min(level-1,4)}")
                add_markdown_runs(p, text, size=Pt({2: 16, 3: 14, 4: 12, 5: 11}.get(level, 11)))
                set_paragraph_format(p, before=8 if level <= 3 else 5, after=3, line=1.02)
                set_keep_with_next(p)
            i += 1
            continue

        if line.strip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = parse_table(block)
            if rows:
                build_table(doc, rows)
            continue

        image = IMAGE_RE.match(line.strip())
        if image:
            insert_image(doc, image.group(2), image.group(1))
            i += 1
            continue

        item_match = ORDERED_RE.match(line.strip())
        if item_match:
            while i < len(lines):
                current = lines[i].strip()
                current_match = ORDERED_RE.match(current)
                if not current_match:
                    break
                num = current_match.group(1)
                item_parts = []
                first_text = current_match.group(2).strip()
                if first_text:
                    item_parts.append(first_text)
                i += 1
                while i < len(lines):
                    nxt = lines[i].rstrip()
                    nxt_stripped = nxt.strip()
                    if not nxt_stripped:
                        break
                    if (
                        HEADING_RE.match(nxt)
                        or nxt_stripped.startswith("|")
                        or IMAGE_RE.match(nxt_stripped)
                        or ORDERED_RE.match(nxt_stripped)
                    ):
                        break
                    item_parts.append(nxt_stripped)
                    i += 1
                p = doc.add_paragraph()
                item_text = " ".join(item_parts).strip()
                add_markdown_runs(p, f"{num}. {item_text}")
                set_paragraph_format(p, before=0, after=2, line=1.12)
            continue

        para_lines = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if HEADING_RE.match(nxt) or nxt.strip().startswith("|") or IMAGE_RE.match(nxt.strip()) or ORDERED_RE.match(nxt.strip()):
                break
            para_lines.append(nxt.strip())
            i += 1
        paragraph = doc.add_paragraph()
        add_markdown_runs(paragraph, " ".join(para_lines))
        set_paragraph_format(paragraph, before=0, after=4, line=1.16)
        set_first_line_indent(paragraph, chars=2.0)

    doc.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_docx()
